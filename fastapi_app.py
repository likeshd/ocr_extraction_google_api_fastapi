from fastapi import FastAPI, UploadFile, File
from fastapi.responses import JSONResponse
import uvicorn
from io import BytesIO
from PIL import Image
import fitz
import tempfile
from docx import Document
import os
import subprocess
from api_functions import pass_ocr_extraction,visa_ocr_extraction,eid_ocr_extraction,dl_ocr_extraction   # Import function from passport_ocr.py


# Initialize FastAPI app
app1 = FastAPI()

# FastAPI route to handle the image upload and OCR extraction
@app1.post("/extract_passport_details")
async def extract_pass_details(image: UploadFile = File(...)):
    try:
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))
            # Get OCR extraction result from the imported function
            data, status_code = await pass_ocr_extraction(image)
            if status_code == 200:
                result = {"data": data, "sts": 200, "msg": "Success"}
                # Return the result as a JSON response
                return JSONResponse(content=result, status_code=status_code)
            elif status_code == 400:
                result = {"msg": "The provided image is not suitable, so the requested details could not be extracted.",
                     "sts": 400}
                return JSONResponse(content=result, status_code=status_code)

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))
            # Get OCR extraction result from the imported function
            data, status_code = await pass_ocr_extraction(image)
            if status_code == 200:
                result = {"data": data, "sts": 200, "msg": "Success"}
                # Return the result as a JSON response
                return JSONResponse(content=result, status_code=status_code)
            elif status_code == 400:
                result = {"msg": "The provided image is not suitable, so the requested details could not be extracted.",
                          "sts": 400}
                return JSONResponse(content=result, status_code=status_code)

        # if file_ext in ['doc', 'docx']:
        #     with tempfile.TemporaryDirectory() as tmpdirname:
        #         input_path = os.path.join(tmpdirname, image.filename)
        #         image.filename.save(input_path)  # Save uploaded file
        #         try:
        #             # Convert DOC to DOCX if necessary
        #             if file_ext == 'doc':
        #                 subprocess.run(
        #                     ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdirname,
        #                      input_path],
        #                     check=True
        #                 )
        #
        #                 # Determine the converted file path
        #                 docx_path = os.path.splitext(input_path)[0] + ".docx"
        #
        #                 # Ensure the converted file exists
        #                 if not os.path.exists(docx_path):
        #                     return JSONResponse(content={"error": "Conversion failed: DOCX file was not created."}, status_code=500)
        #
        #                 # Update input_path to use the converted DOCX file
        #                 input_path = docx_path
        #
        #             # Load the DOCX file
        #             doc = Document(input_path)
        #             for rel in doc.part.rels:
        #                 if "image" in doc.part.rels[rel].target_ref:
        #                     image_data = doc.part.rels[rel].target_part.blob  # Extract image bytes
        #                     image = Image.open(BytesIO(image_data))  # Open as PIL Image
        #                     # Get OCR extraction result from the imported function
        #                     result, status_code = await pass_ocr_extraction(image)
        #                     # Return the result as a JSON response
        #                     return JSONResponse(content=result, status_code=status_code)
        #         except Exception as e:
        #             print(str(e))
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=400)


# FastAPI route to handle the image upload and OCR extraction
@app1.post("/extract_visa_details")
async def extract_pvisa_details(image: UploadFile = File(...)):
    try:
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))
            # Get OCR extraction result from the imported function
            data, status_code = await visa_ocr_extraction(image)
            if status_code == 200:
                result = {"data": data, "sts": 200, "msg": "Success"}
                # Return the result as a JSON response
                return JSONResponse(content=result, status_code=status_code)
            elif status_code == 400:
                result = {"msg": "The provided image is not suitable, so the requested details could not be extracted.",
                          "sts": 400}
                return JSONResponse(content=result, status_code=status_code)

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))
            # Get OCR extraction result from the imported function
            data, status_code = await visa_ocr_extraction(image)
            if status_code == 200:
                result = {"data": data, "sts": 200, "msg": "Success"}
                # Return the result as a JSON response
                return JSONResponse(content=result, status_code=status_code)
            elif status_code == 400:
                result = {"msg": "The provided image is not suitable, so the requested details could not be extracted.",
                          "sts": 400}
                return JSONResponse(content=result, status_code=status_code)

        # if file_ext in ['doc', 'docx']:
        #     with tempfile.TemporaryDirectory() as tmpdirname:
        #         input_path = os.path.join(tmpdirname, image.filename)
        #         image.filename.save(input_path)  # Save uploaded file
        #         try:
        #             # Convert DOC to DOCX if necessary
        #             if file_ext == 'doc':
        #                 subprocess.run(
        #                     ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdirname,
        #                      input_path],
        #                     check=True
        #                 )
        #
        #                 # Determine the converted file path
        #                 docx_path = os.path.splitext(input_path)[0] + ".docx"
        #
        #                 # Ensure the converted file exists
        #                 if not os.path.exists(docx_path):
        #                     return JSONResponse(content={"error": "Conversion failed: DOCX file was not created."}, status_code=500)
        #
        #                 # Update input_path to use the converted DOCX file
        #                 input_path = docx_path
        #
        #             # Load the DOCX file
        #             doc = Document(input_path)
        #             for rel in doc.part.rels:
        #                 if "image" in doc.part.rels[rel].target_ref:
        #                     image_data = doc.part.rels[rel].target_part.blob  # Extract image bytes
        #                     image = Image.open(BytesIO(image_data))  # Open as PIL Image
        #                     # Get OCR extraction result from the imported function
        #                     result, status_code = await visa_ocr_extraction(image)
        #                     # Return the result as a JSON response
        #                     return JSONResponse(content=result, status_code=status_code)
        #         except Exception as e:
        #             print(str(e))
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=400)


# FastAPI route to handle the image upload and OCR extraction
@app1.post("/extract_eid_details")
async def extract_emiratesid_details(image: UploadFile = File(...)):
    try:
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))
            # Get OCR extraction result from the imported function
            data, status_code = await eid_ocr_extraction(image)
            if status_code == 200:
                result = {"data": data, "sts": 200, "msg": "Success"}
                # Return the result as a JSON response
                return JSONResponse(content=result, status_code=status_code)
            elif status_code == 400:
                result = {"msg": "The provided image is not suitable, so the requested details could not be extracted.",
                          "sts": 400}
                return JSONResponse(content=result, status_code=status_code)

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))
            # Get OCR extraction result from the imported function
            data, status_code = await eid_ocr_extraction(image)
            if status_code == 200:
                result = {"data": data, "sts": 200, "msg": "Success"}
                # Return the result as a JSON response
                return JSONResponse(content=result, status_code=status_code)
            elif status_code == 400:
                result = {"msg": "The provided image is not suitable, so the requested details could not be extracted.",
                          "sts": 400}
                return JSONResponse(content=result, status_code=status_code)

        # if file_ext in ['doc', 'docx']:
        #     with tempfile.TemporaryDirectory() as tmpdirname:
        #         input_path = os.path.join(tmpdirname, image.filename)
        #         image.filename.save(input_path)  # Save uploaded file
        #         try:
        #             # Convert DOC to DOCX if necessary
        #             if file_ext == 'doc':
        #                 subprocess.run(
        #                     ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdirname,
        #                      input_path],
        #                     check=True
        #                 )
        #
        #                 # Determine the converted file path
        #                 docx_path = os.path.splitext(input_path)[0] + ".docx"
        #
        #                 # Ensure the converted file exists
        #                 if not os.path.exists(docx_path):
        #                     return JSONResponse(content={"error": "Conversion failed: DOCX file was not created."}, status_code=500)
        #
        #                 # Update input_path to use the converted DOCX file
        #                 input_path = docx_path
        #
        #             # Load the DOCX file
        #             doc = Document(input_path)
        #             for rel in doc.part.rels:
        #                 if "image" in doc.part.rels[rel].target_ref:
        #                     image_data = doc.part.rels[rel].target_part.blob  # Extract image bytes
        #                     image = Image.open(BytesIO(image_data))  # Open as PIL Image
        #                     # Get OCR extraction result from the imported function
        #                     result, status_code = await eid_ocr_extraction(image)
        #                     # Return the result as a JSON response
        #                     return JSONResponse(content=result, status_code=status_code)
        #         except Exception as e:
        #             print(str(e))
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=400)


# FastAPI route to handle the image upload and OCR extraction
@app1.post("/extract_dl_details")
async def extract_driving_license_details(image: UploadFile = File(...)):
    try:
        if  not  image:
            return JSONResponse(content={"error": "No file part in the request"},status_code=400)

        if image.filename == '':
            return JSONResponse(content={"error": "No selected file"},status_code=400)
        file_ext = image.filename.lower().split('.')[-1]

        if file_ext in ['jpg', 'jpeg', 'png']:
            # Read image file
            image_data = await image.read()
            image = Image.open(BytesIO(image_data))
            # Get OCR extraction result from the imported function
            data, status_code = await dl_ocr_extraction(image)
            if status_code == 200:
                result = {"data": data, "sts": 200, "msg": "Success"}
                # Return the result as a JSON response
                return JSONResponse(content=result, status_code=status_code)
            elif status_code == 400:
                result = {"msg": "The provided image is not suitable, so the requested details could not be extracted.",
                          "sts": 400}
                return JSONResponse(content=result, status_code=status_code)

        elif file_ext == 'pdf':
            pdf_bytes = await image.read()
            # Open the PDF file from bytes
            pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
            # img_data_list = []

            # for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(0)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            # Convert Pixmap to PNG bytes
            img_data = pix.tobytes("png")  # Convert directly to PNG bytes
            image = Image.open(BytesIO(img_data))
            # Get OCR extraction result from the imported function
            data, status_code = await dl_ocr_extraction(image)
            if status_code == 200:
                result = {"data": data, "sts": 200, "msg": "Success"}
                # Return the result as a JSON response
                return JSONResponse(content=result, status_code=status_code)
            elif status_code == 400:
                result = {"msg": "The provided image is not suitable, so the requested details could not be extracted.",
                          "sts": 400}
                return JSONResponse(content=result, status_code=status_code)

        # if file_ext in ['doc', 'docx']:
        #     with tempfile.TemporaryDirectory() as tmpdirname:
        #         input_path = os.path.join(tmpdirname, image.filename)
        #         image.filename.save(input_path)  # Save uploaded file
        #         try:
        #             # Convert DOC to DOCX if necessary
        #             if file_ext == 'doc':
        #                 subprocess.run(
        #                     ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdirname,
        #                      input_path],
        #                     check=True
        #                 )
        #
        #                 # Determine the converted file path
        #                 docx_path = os.path.splitext(input_path)[0] + ".docx"
        #
        #                 # Ensure the converted file exists
        #                 if not os.path.exists(docx_path):
        #                     return JSONResponse(content={"error": "Conversion failed: DOCX file was not created."}, status_code=500)
        #
        #                 # Update input_path to use the converted DOCX file
        #                 input_path = docx_path
        #
        #             # Load the DOCX file
        #             doc = Document(input_path)
        #             for rel in doc.part.rels:
        #                 if "image" in doc.part.rels[rel].target_ref:
        #                     image_data = doc.part.rels[rel].target_part.blob  # Extract image bytes
        #                     image = Image.open(BytesIO(image_data))  # Open as PIL Image
        #                     # Get OCR extraction result from the imported function
        #                     result, status_code = await dl_ocr_extraction(image)
        #                     # Return the result as a JSON response
        #                     return JSONResponse(content=result, status_code=status_code)
        #         except Exception as e:
        #             print(str(e))
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=400)


# Testing the FastAPI app locally
if __name__ == "__main__":
    # Run the FastAPI application using uvicorn
    # uvicorn.run(app, host="127.0.0.1", port=8000)
    uvicorn.run(app1, host="0.0.0.0", port=8000)